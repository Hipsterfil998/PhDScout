"""CV parser: reads PDF / DOCX / TXT and returns a structured research profile."""

from __future__ import annotations

from pathlib import Path
from typing import Any, TypedDict

from agent.llm_client import LLMClient
from agent.utils import parse_json


# ---------------------------------------------------------------------------
# Type definition
# ---------------------------------------------------------------------------

class CVProfile(TypedDict, total=False):
    name: str
    contact: dict                  # email, phone, linkedin, github, website
    summary: str
    education: list[dict]          # degree, institution, field, year, thesis_topic
    experience: list[dict]         # title, institution, dates, description
    research_interests: list[str]
    publications: list[dict]       # title, venue, year, authors
    skills: dict                   # programming, tools, languages, lab_techniques
    awards: list[str]
    languages: list[dict]          # language, level
    references: list[dict]


# ---------------------------------------------------------------------------
# Raw text extraction (module-level — pure I/O, no LLM)
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(path: Path) -> str:
    import pdfplumber  # type: ignore
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _extract_text_from_docx(path: Path) -> str:
    from docx import Document  # type: ignore
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_raw_text(cv_path: str | Path) -> str:
    """Extract raw text from a CV file (.pdf, .docx, .txt)."""
    path = Path(cv_path)
    if not path.exists():
        raise FileNotFoundError(f"CV file not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_from_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _extract_text_from_docx(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported CV format '{suffix}'. Use .pdf, .docx, or .txt.")


# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

_SYSTEM = (
    "You are an expert academic CV parser. "
    "Extract ALL information from the CV text into structured JSON. "
    "Pay special attention to research-specific fields: thesis topics, "
    "publications, research interests, and lab/technical skills. "
    "Respond only with valid JSON — no markdown fences, no commentary."
)

_PROMPT = """Extract ALL information from the following CV and return a single JSON object.

Expected structure (use null for missing scalars, [] for missing lists):
{{
  "name": "Full Name",
  "contact": {{"email": "...", "phone": "...", "linkedin": "...", "github": "...", "website": "..."}},
  "summary": "Brief research summary",
  "education": [{{"degree": "PhD", "institution": "MIT", "field": "CS", "year": "2021", "thesis_topic": "..."}}],
  "experience": [{{"title": "Research Assistant", "institution": "ETH Zurich", "dates": "2019-2021", "description": "..."}}],
  "research_interests": ["machine learning", "NLP"],
  "publications": [{{"title": "...", "venue": "NeurIPS 2022", "year": "2022", "authors": "..."}}],
  "skills": {{"programming": ["Python"], "tools": ["PyTorch"], "languages": ["LaTeX"], "lab_techniques": []}},
  "awards": ["Best Paper Award NeurIPS 2022"],
  "languages": [{{"language": "English", "level": "Native"}}],
  "references": [{{"name": "Prof. Jane Smith", "title": "Full Professor", "institution": "MIT"}}]
}}

Do NOT invent information — extract only what is present.

CV TEXT:
---
{cv_text}
---"""


# ---------------------------------------------------------------------------
# CVParser class
# ---------------------------------------------------------------------------

class CVParser:
    """Parses CV files into structured CVProfile dicts using an LLM."""

    def __init__(self, llm: LLMClient) -> None:
        self.llm = llm

    def parse(self, cv_path: str | Path) -> CVProfile:
        """Parse a CV file and return a structured CVProfile."""
        raw_text = extract_raw_text(cv_path)
        if not raw_text.strip():
            raise ValueError("Could not extract any text from the CV file.")

        prompt = _PROMPT.format(cv_text=raw_text[:8000])
        raw_json = self.llm.generate(system=_SYSTEM, user=prompt, json_mode=True)

        result = parse_json(raw_json)
        if result is None:
            return {"name": "Unknown", "summary": raw_json[:500]}
        return result

    def summarize(self, profile: CVProfile) -> str:
        """Build a compact text summary of a CVProfile for use in LLM prompts."""
        lines: list[str] = []

        if profile.get("name"):
            lines.append(f"Name: {profile['name']}")

        contact: dict = profile.get("contact") or {}
        if contact.get("email"):
            lines.append(f"Email: {contact['email']}")

        if profile.get("summary"):
            lines.append(f"Summary: {profile['summary']}")

        research = profile.get("research_interests") or []
        if research:
            lines.append(f"Research interests: {', '.join(research[:10])}")

        for e in (profile.get("education") or [])[:3]:
            thesis = f" — Thesis: {e['thesis_topic']}" if e.get("thesis_topic") else ""
            lines.append(
                f"Education: {e.get('degree', '')} in {e.get('field', '')} "
                f"from {e.get('institution', '')} ({e.get('year', '')}){thesis}"
            )

        pubs = profile.get("publications") or []
        if pubs:
            lines.append(f"Publications ({len(pubs)}):")
            for p in pubs[:5]:
                lines.append(f"  - \"{p.get('title', '')}\" — {p.get('venue', '')} {p.get('year', '')}")

        for e in (profile.get("experience") or [])[:4]:
            lines.append(f"Experience: {e.get('title', '')} at {e.get('institution', '')} ({e.get('dates', '')})")

        skills: dict = profile.get("skills") or {}
        all_skills = (skills.get("programming") or []) + (skills.get("tools") or [])
        if all_skills:
            lines.append(f"Technical skills: {', '.join(all_skills[:20])}")
        lab = skills.get("lab_techniques") or []
        if lab:
            lines.append(f"Lab techniques: {', '.join(lab[:10])}")

        awards = profile.get("awards") or []
        if awards:
            lines.append(f"Awards: {'; '.join(awards[:5])}")

        langs = profile.get("languages") or []
        if langs:
            lines.append(
                "Languages: " + ", ".join(
                    f"{la.get('language', '')} ({la.get('level', '')})" for la in langs
                )
            )

        return "\n".join(lines)
